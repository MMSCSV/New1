import tkinter as tk
from tkinter import Toplevel, Canvas
from PIL import ImageGrab
import os

counter_1=1.1
counter_2=0
counter_3=1
x1,y1,x2,y2=0,0,0,0
fl=0;

save_folder="scr"

if not os.path.exists(save_folder):
    os.makedirs(save_folder)
va=0;

def Generate_Report():
    from docx import Document
    from docx.shared import Inches
    import os
    from os import listdir

    document = Document()

    folder_dir=r"C:\Users\11002733\source\repos\PythonApplication3\scr"
    document.add_heading('Objective evidence', 0)
    for images in os.listdir(folder_dir):
      if images[2]=='1':
        document.add_heading("Test Case"+images[0])
        document.add_paragraph("step"+images[2])
      elif images[2]!='1' and images[2]!='e':
        document.add_paragraph("step"+images[2])
      elif images[2]=='e':
        document.add_heading("Expected Result")
    
      Rep="C:\\Users\\11002733\\source\\repos\\PythonApplication3\\scr\\"+images
      document.add_picture(Rep, width=Inches(5))

    document.add_page_break()

    document.save('demo9.docx')

def set_capture_area():
    global x1, y1, x2, y2

    def on_button_press(event):
        global x1, y1
        x1, y1 = event.x, event.y
        canvas.delete("rectangle")

    def on_mouse_drag(event):
        global x2, y2
        x2, y2 = event.x, event.y
        canvas.delete("rectangle")
        canvas.create_rectangle(x1, y1, x2, y2, outline='red', width=2, tag="rectangle")

    def on_button_release(event):
        global x2, y2
        x2, y2 = event.x, event.y
        capture_window.destroy()
        print(f"Capture area set to: {(x1, y1, x2, y2)}")

    capture_window = Toplevel(root)
    capture_window.attributes('-fullscreen', True)
    capture_window.attributes('-alpha', 0.3)
    canvas = Canvas(capture_window, cursor="cross")
    canvas.pack(fill=tk.BOTH, expand=True)
    canvas.bind("<ButtonPress-1>", on_button_press)
    canvas.bind("<B1-Motion>", on_mouse_drag)
    canvas.bind("<ButtonRelease-1>", on_button_release)


def take_snapshot_1():
    global counter_1
   
    global fl;
    fl=round(counter_1,3)
    filename =f"{fl}.png"
    counter_1 +=1
    take_snapshot(filename)
    
def take_snapshot_2():
    global counter_2
    global va;
    va+=0.1
    counter_2 =fl+va
    al=round(counter_2,2)
    filename =f"{al}.png"
    take_snapshot(filename)
    

def take_snapshot(filename):
    global x1, y1, x2, y2
    if x1 == y1 == x2 == y2 == 0:
        print("Capture area is not set. Please set the capture area first.")
        return
   
    # Capture the specified area
    snapshot = ImageGrab.grab(bbox=(x1, y1, x2, y2))
   
    # Save the snapshot to the specified folder
    filepath = os.path.join(save_folder, filename)
    snapshot.save(filepath)
    print(f"Snapshot saved as {filepath}")

# Create the main window
root = tk.Tk()
root.title("Snapshot Taker")

# Create a button to set the capture area
btn_set_area = tk.Button(root, text="Set Capture Area", command=set_capture_area)
btn_set_area.pack(pady=5)

# Create buttons to take snapshots
btn_snapshot_1 = tk.Button(root, text="First Step", command=take_snapshot_1)
btn_snapshot_1.pack(pady=5)

btn_snapshot_2 = tk.Button(root, text="Other Steps", command=take_snapshot_2)
btn_snapshot_2.pack(pady=5)

btn_Generate_Report = tk.Button(root, text="Generate_Report", command=Generate_Report)
btn_Generate_Report.pack(pady=5)


# Run the application
root.mainloop()
